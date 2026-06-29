"""
Skillsbucket — Proposal Assistant (Streamlit Frontend)

This file provides a clean, modern Streamlit UI for the Skillsbucket RAG
proposal assistant. It keeps all existing functionality (chat, proposal
generation, docx download) while fixing Streamlit API usage, CSS issues,
and improving styling and UI components.

Requirements preserved: RAG retrieval via FastAPI, proposal generation,
DOCX export, client intake form, chat history.
"""

from datetime import date
import io
import re
import html
from typing import List, Optional

import requests
import streamlit as st

# Page configuration (must be called before other Streamlit UI calls)
st.set_page_config(
    page_title="Skillsbucket — Proposal Assistant",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)

API_BASE = "http://localhost:8000"


# =========================
# CUSTOM CSS (Premium Dark Mode)
# =========================
st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

/* Apply font family via inheritance to prevent breaking icon fonts */
html, body, [data-testid="stAppViewContainer"] {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
}

/* App Background (Dark Mode) */
.stApp {
    background-color: #0B0F19 !important;
    color: #F8FAFC !important;
}

/* Sidebar styling - Dark Mode */
section[data-testid="stSidebar"] {
    background-color: #0F172A !important;
    color: #F8FAFC !important;
    border-right: 1px solid #1E293B;
}

section[data-testid="stSidebar"] label {
    color: #94A3B8 !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
}

/* Inputs in Sidebar: Dark Mode background with white text */
section[data-testid="stSidebar"] .stTextInput>div>div>input,
section[data-testid="stSidebar"] .stTextArea>div>div>textarea,
section[data-testid="stSidebar"] .stNumberInput>div>div>input,
section[data-testid="stSidebar"] [data-baseweb="select"] {
    background-color: #1E293B !important;
    color: #F8FAFC !important;
    border: 1px solid #334155 !important;
    border-radius: 8px !important;
}

/* Make sure the selected option value inside the selectbox is visible (white text) */
section[data-testid="stSidebar"] [data-baseweb="select"] [data-testid="stSelectboxSelectedValue"],
section[data-testid="stSidebar"] [data-baseweb="select"] span,
section[data-testid="stSidebar"] [data-baseweb="select"] div {
    color: #F8FAFC !important;
}

/* Dropdown popover list styling - dark background, white text, blue hover */
div[data-baseweb="popover"] ul,
div[data-baseweb="menu"],
[role="listbox"] {
    background-color: #1E293B !important;
    border: 1px solid #334155 !important;
    border-radius: 8px !important;
    box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.3), 0 4px 6px -2px rgba(0, 0, 0, 0.2) !important;
}

div[data-baseweb="popover"] li,
div[data-baseweb="menu"] div,
div[data-baseweb="menu"] li,
[role="option"] {
    color: #F8FAFC !important;
    background-color: #1E293B !important;
    padding: 10px 14px !important;
    font-size: 0.9rem !important;
    transition: all 0.15s ease !important;
}

div[data-baseweb="popover"] li:hover,
div[data-baseweb="menu"] div:hover,
div[data-baseweb="menu"] li:hover,
[role="option"]:hover,
[aria-selected="true"] {
    background-color: #2563EB !important;
    color: #FFFFFF !important;
}

/* Tabs customization for Dark Mode */
.stTabs [data-baseweb="tab"] {
    color: #94A3B8 !important;
}
.stTabs [data-baseweb="tab"][aria-selected="true"] {
    color: #FFFFFF !important;
    border-bottom-color: #2563EB !important;
}

/* Header styling */
.app-header {
    background: linear-gradient(135deg, #1E293B 0%, #1E3A8A 50%, #2563EB 100%);
    padding: 28px 32px;
    border-radius: 16px;
    color: white;
    margin-bottom: 24px;
    box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2);
    border: 1px solid #334155;
}
.app-header h1 {
    margin: 0;
    font-size: 1.85rem;
    font-weight: 700;
    letter-spacing: -0.025em;
    color: #FFFFFF !important;
}
.app-header p {
    margin: 8px 0 0 0;
    color: #93C5FD !important;
    font-size: 0.95rem;
}

/* Chat bubble designs for Dark Mode */
.chat-user {
    background: linear-gradient(135deg, #2563EB 0%, #1D4ED8 100%);
    color: white;
    padding: 14px 18px;
    border-radius: 16px 16px 2px 16px;
    margin: 12px 0 12px auto;
    max-width: 75%;
    box-shadow: 0 4px 15px rgba(37, 99, 235, 0.2);
    font-size: 0.95rem;
    line-height: 1.5;
}
.chat-assistant {
    background: #1E293B;
    color: #F8FAFC;
    padding: 14px 18px;
    border-radius: 16px 16px 16px 2px;
    margin: 12px auto 12px 0;
    max-width: 75%;
    box-shadow: 0 4px 15px rgba(0, 0, 0, 0.15);
    border: 1px solid #334155;
    font-size: 0.95rem;
    line-height: 1.5;
}
.source-tag {
    font-size: 0.8rem;
    color: #94A3B8;
    background-color: #1E293B;
    padding: 6px 12px;
    border-radius: 20px;
    display: inline-block;
    margin-top: 6px;
    border: 1px solid #334155;
    font-weight: 500;
}

/* Modern Document Container for Proposal - Dark Mode */
.proposal-container {
    background: #1E293B;
    padding: 40px 48px;
    border-radius: 16px;
    border: 1px solid #334155;
    box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.3);
    margin-top: 16px;
    line-height: 1.7;
    color: #E2E8F0;
}

/* Modern tables inside Dark Mode */
.modern-table {
    width: 100%;
    border-collapse: separate;
    border-spacing: 0;
    color: #E2E8F0;
    background: #1E293B;
    border-radius: 12px;
    overflow: hidden;
    border: 1px solid #334155;
    margin: 16px 0;
}
.modern-table th {
    background: #0F172A;
    color: white;
    padding: 12px 16px;
    font-weight: 600;
    text-align: left;
    font-size: 0.9rem;
    border-bottom: 2px solid #334155;
}
.modern-table td {
    padding: 12px 16px;
    font-size: 0.9rem;
    border-bottom: 1px solid #334155;
}
.modern-table tr:last-child td {
    border-bottom: none;
}
.modern-table tr:nth-child(even) {
    background-color: #111827;
}

/* High-contrast warning box - Dark Mode optimized */
.warning-box {
    background-color: #2D251E !important;
    color: #FBBF24 !important;
    padding: 16px 20px !important;
    border-radius: 12px !important;
    border-left: 5px solid #F59E0B !important;
    margin: 16px 0 !important;
    font-weight: 500 !important;
    font-size: 0.95rem !important;
    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.2) !important;
    display: flex;
    align-items: center;
    gap: 12px;
}

/* Global button styles to align with professional UI */
.stButton>button, .stDownloadButton>button {
    background-color: #2563EB !important;
    color: white !important;
    border-radius: 8px !important;
    border: none !important;
    padding: 8px 16px !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    transition: all 0.25s ease !important;
    box-shadow: 0 2px 4px rgba(37, 99, 235, 0.2) !important;
}
.stButton>button:hover, .stDownloadButton>button:hover {
    background-color: #1D4ED8 !important;
    box-shadow: 0 4px 8px rgba(37, 99, 235, 0.3) !important;
    transform: translateY(-1px);
}
.stButton>button:active, .stDownloadButton>button:active {
    transform: translateY(0);
}

/* Responsive design */
@media (max-width: 768px) {
    .chat-user, .chat-assistant {
        max-width: 90%;
    }
    .proposal-container {
        padding: 24px;
    }
}
</style>
""",
    unsafe_allow_html=True,
)


# =========================
# Helper functions
# =========================

def check_api_health(timeout: float = 3.0) -> bool:
    try:
        r = requests.get(f"{API_BASE}/health", timeout=timeout)
        return r.status_code == 200
    except Exception:
        return False


def call_chat_api(message: str, client_context: dict, timeout: float = 60.0) -> dict:
    payload = {"message": message, "client_context": client_context}
    r = requests.post(f"{API_BASE}/chat", json=payload, timeout=timeout)
    r.raise_for_status()
    return r.json()


def call_proposal_api(client_context: dict, timeout: float = 120.0) -> dict:
    payload = {"client_context": client_context}
    r = requests.post(f"{API_BASE}/proposal", json=payload, timeout=timeout)
    r.raise_for_status()
    return r.json()


def show_warning_box(message: str) -> None:
    st.markdown(f'<div class="warning-box"><span>⚠️</span><span>{html.escape(message)}</span></div>', unsafe_allow_html=True)


def render_chat_message(role: str, content: str, sources: int = 0, cats: Optional[List[str]] = None) -> None:
    if role == "user":
        st.markdown(f'<div class="chat-user">{html.escape(content)}</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="chat-assistant">{html.escape(content)}</div>', unsafe_allow_html=True)
        if sources:
            cats_str = ", ".join(cats or [])
            st.markdown(f'<div class="source-tag">📚 {sources} sources · {cats_str}</div>', unsafe_allow_html=True)


def markdown_to_docx(md_text: str) -> bytes:
    # Lightweight MD -> DOCX converter preserving headings, bullets and basic bold
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.core_properties.author = "Skillsbucket"
    para_buf = []
    for line in md_text.splitlines():
        line = line.rstrip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.strip() in ("", "---"):
            doc.add_paragraph("")
        else:
            para = doc.add_paragraph()
            # simple bold parsing for **bold**
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# =========================
# Session state defaults
# =========================
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "proposal_output" not in st.session_state:
    st.session_state.proposal_output = None

if "proposal_meta" not in st.session_state:
    st.session_state.proposal_meta = {}


# =========================
# Sidebar — Client intake
# =========================
with st.sidebar:
    st.markdown(
        """
    <div style="text-align:center; padding: 8px 0 18px 0;">
        <div style="font-size: 2.5rem;">🎓</div>
        <div style="color: white; font-size: 1.15rem; font-weight: 700; margin-top: 8px;">Skillsbucket</div>
        <div style="color: #94A3B8; font-size: 0.8rem; margin-top: 4px;">People · Purpose · Progress</div>
    </div>
    """,
        unsafe_allow_html=True,
    )

    st.markdown("---")
    st.markdown("### 📋 Client Details")

    client_name = st.text_input("Client Name", value="Rajesh Kumar", placeholder="Enter contact person name")
    company_name = st.text_input("Company / Organization", value="HDFC Bank", placeholder="Enter company name")

    industry = st.selectbox(
        "Target Industry",
        options=[
            "Banking and Financial Services",
            "Information Technology",
            "Manufacturing",
            "Healthcare and Pharmaceuticals",
            "Retail and FMCG",
            "Infrastructure and Real Estate",
            "Telecom",
            "Hospitality",
            "Other",
        ],
    )

    training_course = st.selectbox(
        "Training Program",
        options=[
            "Leadership Skills for First Time Managers",
            "Leadership Skills for Senior Leaders",
            "Mindful Leadership Through Emotional Intelligence",
            "Leading in a VUCA World",
            "Managing Millennials",
            "Impactful Business Communication",
            "Business Writing Skills",
            "Business Storytelling",
            "Business Etiquette",
            "Negotiation Skills",
            "Consultative Selling",
            "Customer Centricity",
            "Influence to Win",
            "Sales Channel Management",
            "Stress Management",
            "Time Management",
            "Problem Solving and Creativity",
            "Accountability and Ownership",
            "Assertive Communication",
            "Conflict Management",
            "Team Goal Setting",
            "Interpersonal Communication",
            "Outbound Team Building",
            "Campus to Corporate",
            "Design Thinking",
            "Diversity and Inclusion",
            "Executive Coaching",
            "Train the Trainer",
            "Competency Mapping",
            "Custom Program",
        ],
    )

    delivery_mode = st.selectbox(
        "Delivery Mode",
        options=[
            "Instructor-Led Training (ILT) - Offline",
            "Virtual Instructor-Led Training (VILT) - Online",
            "Blended Learning (ILT + Digital)",
        ],
    )

    audience = st.selectbox(
        "Target Audience",
        options=[
            "Mid-level managers",
            "Senior leadership / CXO",
            "First-time managers / Team leads",
            "Individual contributors",
            "Sales teams",
            "Fresh graduates / Campus hires",
            "Women professionals",
            "Cross-functional teams",
        ],
    )

    participants = st.number_input("Number of Participants", min_value=5, max_value=200, value=25, step=5)

    duration = st.selectbox(
        "Program Duration",
        options=["Half Day (4 hours)", "1 Day", "2 Days", "3 Days", "4 Weeks (Modular)", "Custom"],
    )

    training_dates = st.text_input("Preferred Training Dates", value="To be confirmed", placeholder="e.g., July 15-16, 2025")

    special_requirements = st.text_area(
        "Special Requirements", placeholder="Any customization, language preferences, industry case studies...", height=80
    )

    st.markdown("---")

    api_ok = check_api_health()
    if api_ok:
        st.markdown('<p style="color:#10B981; font-weight:700;">🟢 API Connected</p>', unsafe_allow_html=True)
    else:
        st.markdown('<p style="color:#EF4444; font-weight:700;">🔴 API Offline — run: uvicorn app_api:app --reload --port 8000</p>', unsafe_allow_html=True)

    if st.button("🗑️ Clear Chat"):
        st.session_state.chat_history = []
        st.rerun()


# Client context assembled from sidebar inputs
client_ctx = {
    "client_name": client_name,
    "company_name": company_name,
    "industry": industry,
    "training_course": training_course,
    "delivery_mode": delivery_mode,
    "participants": str(participants),
    "duration": duration,
    "training_dates": training_dates,
    "special_requirements": special_requirements or "None",
}


# =========================
# Main layout
# =========================
st.markdown(
    """
<div class="app-header">
    <h1>🎓 Skillsbucket Proposal Assistant</h1>
    <p>AI-powered corporate training proposal generator · Powered by Verified Knowledge Base</p>
</div>
""",
    unsafe_allow_html=True,
)

tab_chat, tab_proposal = st.tabs(["💬  Chat Assistant", "📄  Generate Proposal"])


# -------------------------
# TAB 1: Chat Assistant
# -------------------------
with tab_chat:
    st.markdown("#### Ask anything about Skillsbucket programs, pricing, or services")
    st.caption("The assistant only answers from the verified knowledge base. Off-topic questions are politely refused.")

    col1, col2, col3, col4 = st.columns(4)
    suggestions = [
        "What leadership programs do you offer?",
        "What is the pricing for ILT programs?",
        "Tell me about Skillsbucket's experience",
        "What are the payment terms?",
    ]
    if col1.button(suggestions[0]):
        st.session_state["quick_input"] = suggestions[0]
    if col2.button(suggestions[1]):
        st.session_state["quick_input"] = suggestions[1]
    if col3.button(suggestions[2]):
        st.session_state["quick_input"] = suggestions[2]
    if col4.button(suggestions[3]):
        st.session_state["quick_input"] = suggestions[3]

    st.markdown("---")

    with st.container():
        if not st.session_state.chat_history:
            st.markdown(
                """
            <div style="text-align:center; padding: 40px; color: #94A3B8;">
                <div style="font-size: 2.5rem;">💬</div>
                <div style="font-size: 1rem; margin-top: 8px;">Start a conversation about Skillsbucket programs and services</div>
            </div>
            """,
                unsafe_allow_html=True,
            )
        else:
            for msg in st.session_state.chat_history:
                render_chat_message(
                    role=msg.get("role", "assistant"),
                    content=msg.get("content", ""),
                    sources=msg.get("sources", 0),
                    cats=msg.get("categories", []),
                )

    default_input = st.session_state.pop("quick_input", "")
    user_input = st.chat_input(placeholder="Ask about courses, pricing, methodology, about Skillsbucket...")

    if user_input or default_input:
        message = user_input or default_input
        st.session_state.chat_history.append({"role": "user", "content": message})

        if not api_ok:
            show_warning_box("API is not running. Please start the FastAPI server: uvicorn app_api:app --reload --port 8000")
        else:
            with st.spinner("Searching knowledge base..."):
                try:
                    result = call_chat_api(message, client_ctx)
                    st.session_state.chat_history.append(
                        {
                            "role": "assistant",
                            "content": result.get("reply", ""),
                            "sources": result.get("sources_matched", 0),
                            "categories": result.get("categories", []),
                        }
                    )
                except Exception as e:
                    st.session_state.chat_history.append({"role": "assistant", "content": f"⚠️ Error: {str(e)}", "sources": 0, "categories": []})
        st.rerun()


# -------------------------
# TAB 2: Proposal Generation
# -------------------------
with tab_proposal:
    st.markdown("#### Generate a Complete Corporate Training Proposal")
    st.caption(
        "Uses all sidebar inputs to retrieve relevant course, pricing, T&C, and company information from the knowledge base and generate a full proposal."
    )

    with st.expander("📋 Current Client Intake (from sidebar)", expanded=True):
        st.markdown(
            f"""
<table class="modern-table">
  <thead>
    <tr>
      <th style="width:35%;">Field</th>
      <th>Value</th>
    </tr>
  </thead>
  <tbody>
    <tr><td style="font-weight:600;">Client Name</td><td>{client_name}</td></tr>
    <tr><td style="font-weight:600;">Company</td><td>{company_name}</td></tr>
    <tr><td style="font-weight:600;">Industry</td><td>{industry}</td></tr>
    <tr><td style="font-weight:600;">Program</td><td>{training_course}</td></tr>
    <tr><td style="font-weight:600;">Delivery Mode</td><td>{delivery_mode}</td></tr>
    <tr><td style="font-weight:600;">Audience</td><td>{audience}</td></tr>
    <tr><td style="font-weight:600;">Participants</td><td>{participants}</td></tr>
    <tr><td style="font-weight:600;">Duration</td><td>{duration}</td></tr>
    <tr><td style="font-weight:600;">Training Dates</td><td>{training_dates}</td></tr>
  </tbody>
</table>
""",
            unsafe_allow_html=True,
        )

    # Generate button (no unsupported kwargs)
    generate_btn = st.button("🚀 Generate Full Proposal", disabled=not api_ok)

    if not api_ok:
        show_warning_box("Start the API server first: uvicorn app_api:app --reload --port 8000")

    if generate_btn and api_ok:
        with st.spinner(f"Generating proposal for {company_name}... (this may take 20-40 seconds)"):
            try:
                result = call_proposal_api(client_ctx)
                st.session_state.proposal_output = result.get("proposal_markdown", "")
                st.session_state.proposal_meta = {
                    "sources": result.get("sources_matched", 0),
                    "categories": result.get("categories", []),
                }
                st.success(
                    f"✅ Proposal generated — {st.session_state.proposal_meta['sources']} knowledge base sources used across {len(st.session_state.proposal_meta['categories'])} categories."
                )
            except Exception as e:
                st.error(f"Generation failed: {str(e)}")

    if st.session_state.proposal_output:
        meta = st.session_state.proposal_meta

        btn_col1, btn_col2, btn_col3 = st.columns([2, 1, 1])

        with btn_col1:
            docx_bytes = markdown_to_docx(st.session_state.proposal_output)
            filename = f"Skillsbucket_Proposal_{company_name.replace(' ', '_')}_{date.today()}.docx"
            st.download_button(
                label="⬇️ Download Proposal (.docx)",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        with btn_col2:
            if st.button("🔄 Regenerate"):
                st.session_state.proposal_output = None
                st.rerun()

        with btn_col3:
            if st.button("🗑️ Clear"):
                st.session_state.proposal_output = None
                st.rerun()

        st.markdown("---")

        st.markdown('<div class="proposal-container">', unsafe_allow_html=True)
        st.markdown(st.session_state.proposal_output)
        st.markdown('</div>', unsafe_allow_html=True)

        with st.expander("📚 Knowledge Base Sources Used"):
            st.write(f"**Total chunks retrieved:** {meta.get('sources', 0)}")
            st.write("**Categories covered:**")
            for cat in meta.get("categories", []):
                st.markdown(f"- {cat}")
